from flask import Flask, render_template, request, redirect, url_for, flash, send_file, session, jsonify
from functools import wraps
import pandas as pd
import os
from datetime import datetime
import matplotlib
matplotlib.use('Agg')  # Use non-interactive backend
import matplotlib.pyplot as plt
import io
import base64
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from database import (
    init_database, insert_pidum_data, insert_pidsus_data,
    get_all_pidum_data, get_all_pidsus_data,
    get_pidum_data_for_export, get_pidsus_data_for_export,
    get_database_stats, get_pidum_report_data,
    delete_all_pidum_data, delete_pidum_item,
    authenticate_user
)
from import_helper import process_import_file, get_jenis_perkara_suggestions, prepare_import_data
from import_pra_penuntutan_helper import (
    process_pra_penuntutan_import_file, 
    get_jenis_perkara_suggestions_pra_penuntutan,
    prepare_pra_penuntutan_data_for_db,
    allowed_file_pra_penuntutan
)
from import_upaya_hukum_helper import (
    process_upaya_hukum_import_file,
    get_jenis_perkara_suggestions_upaya_hukum,
    prepare_upaya_hukum_data_for_db,
    allowed_file_upaya_hukum
)

def generate_pidum_chart(report_data):
    """Generate chart data for PIDUM report"""
    # Create figure
    plt.figure(figsize=(12, 6))
    
    # Prepare data for chart - show all jenis perkara including those with value 0
    jenis_perkara = [item['jenis_perkara'] for item in report_data]
    jumlah_data = [item['JUMLAH'] for item in report_data]
    
    if not jenis_perkara:  # If no data at all, create empty chart with all categories
        jenis_perkara = ['NARKOBA', 'PERKARA ANAK', 'KESUSILAAN', 'JUDI', 'KDRT', 'OHARDA', 'PERKARA LAINNYA']
        jumlah_data = [0] * 7
    
    # Create bar chart
    bars = plt.bar(range(len(jenis_perkara)), jumlah_data, color='#4472C4')
    
    # Customize chart
    plt.title('Grafik PIDUM', fontsize=14, fontweight='bold', pad=20)
    plt.xlabel('Jenis Perkara', fontsize=12)
    plt.ylabel('Jumlah', fontsize=12)
    
    # Set x-axis labels
    plt.xticks(range(len(jenis_perkara)), jenis_perkara, rotation=45, ha='right')
    
    # Add value labels on bars (show all values, including 0)
    for i, (bar, value) in enumerate(zip(bars, jumlah_data)):
        plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.1, 
                str(value), ha='center', va='bottom', fontweight='bold')
    
    # Set y-axis to start from 0 and add some padding
    max_value = max(jumlah_data) if jumlah_data and max(jumlah_data) > 0 else 5
    plt.ylim(0, max_value + 1)
    
    # Add grid for better readability
    plt.grid(axis='y', alpha=0.3)
    
    # Adjust layout to prevent label cutoff
    plt.tight_layout()
    
    # Convert to base64 string
    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
    img_buffer.seek(0)
    chart_base64 = base64.b64encode(img_buffer.getvalue()).decode()
    plt.close()
    
    return chart_base64

app = Flask(__name__)
app.secret_key = 'your_secret_key_here'
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024  # 10MB max file size

# Initialize database on startup
init_database()

# Custom Jinja2 filter for summing numeric strings
@app.template_filter('sum_numeric')
def sum_numeric(data_list, attribute):
    """Sum numeric values from a list of dictionaries, converting strings to integers"""
    try:
        total = 0
        for item in data_list:
            value = item.get(attribute, 0)
            # Convert to int, handle both string and int types
            if isinstance(value, str):
                total += int(value) if value.isdigit() else 0
            else:
                total += int(value)
        return total
    except (ValueError, TypeError):
        return 0

# Login decorator
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        
        user = authenticate_user(username, password)
        
        if user:
            session['user_id'] = user['id']
            session['username'] = user['username']
            flash('Login berhasil!', 'success')
            return redirect(url_for('index'))
        else:
            flash('Username atau password salah!', 'error')
    
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.clear()
    flash('Anda telah keluar dari sistem.', 'success')
    return redirect(url_for('login'))

@app.route('/')
@login_required
def index():
    stats = get_database_stats()
    return render_template('index.html', stats=stats)

@app.route('/pilihan_laporan')
@login_required
def pilihan_laporan():
    return render_template('pilihan_laporan.html')

@app.route('/laporan_pidum_bulanan')
@login_required
def laporan_pidum_bulanan():
    # Get filter parameters
    tahun = request.args.get('tahun', type=int, default=2025)
    bulan = request.args.get('bulan', type=int)
    
    from datetime import datetime
    import sqlite3
    from collections import defaultdict
    
    # Get database connection
    conn = sqlite3.connect('db/kejaksaan.db')
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    # Build query with year and optional month filter
    where_conditions = ["strftime('%Y', tanggal) = ?"]
    params = [str(tahun)]
    
    if bulan:
        where_conditions.append("strftime('%m', tanggal) = ?")
        params.append(f"{bulan:02d}")
    
    where_clause = "WHERE " + " AND ".join(where_conditions)
    
    # Get data from pidum_data table
    query = f"""
    SELECT periode, jenis_perkara, tanggal, tahapan_penanganan,
           strftime('%m', tanggal) as bulan_num,
           CASE strftime('%m', tanggal)
               WHEN '01' THEN 'Januari'
               WHEN '02' THEN 'Februari'
               WHEN '03' THEN 'Maret'
               WHEN '04' THEN 'April'
               WHEN '05' THEN 'Mei'
               WHEN '06' THEN 'Juni'
               WHEN '07' THEN 'Juli'
               WHEN '08' THEN 'Agustus'
               WHEN '09' THEN 'September'
               WHEN '10' THEN 'Oktober'
               WHEN '11' THEN 'November'
               WHEN '12' THEN 'Desember'
           END as bulan_nama
    FROM pidum_data {where_clause}
    ORDER BY bulan_num, jenis_perkara
    """
    
    cursor.execute(query, params)
    rows = cursor.fetchall()
    
    # Get available months for filter dropdown
    month_query = f"""
    SELECT DISTINCT strftime('%m', tanggal) as bulan_num,
           CASE strftime('%m', tanggal)
               WHEN '01' THEN 'Januari'
               WHEN '02' THEN 'Februari'
               WHEN '03' THEN 'Maret'
               WHEN '04' THEN 'April'
               WHEN '05' THEN 'Mei'
               WHEN '06' THEN 'Juni'
               WHEN '07' THEN 'Juli'
               WHEN '08' THEN 'Agustus'
               WHEN '09' THEN 'September'
               WHEN '10' THEN 'Oktober'
               WHEN '11' THEN 'November'
               WHEN '12' THEN 'Desember'
           END as bulan_nama
    FROM pidum_data
    WHERE strftime('%Y', tanggal) = ?
    ORDER BY bulan_num
    """
    
    cursor.execute(month_query, [str(tahun)])
    available_months = cursor.fetchall()
    
    conn.close()
    
    # Define all categories that should always appear
    predefined_categories = ['NARKOBA', 'PERKARA ANAK', 'KESUSILAAN', 'JUDI', 'KDRT', 'OHARDA', 'PERKARA LAINNYA']
    
    # Get all months that should be displayed
    month_names = {
        1: 'Januari', 2: 'Februari', 3: 'Maret', 4: 'April',
        5: 'Mei', 6: 'Juni', 7: 'Juli', 8: 'Agustus',
        9: 'September', 10: 'Oktober', 11: 'November', 12: 'Desember'
    }
    
    # Determine which months to display
    if bulan:
        display_months = [month_names.get(bulan, 'Januari')]
    else:
        # Show all months that have data or all 12 months if no specific month is selected
        months_with_data = set()
        for month in available_months:
            months_with_data.add(month['bulan_nama'])
        
        # If no specific month is selected, show all months that have data
        display_months = list(months_with_data) if months_with_data else list(month_names.values())
    
    # Process data for report
    data_summary = defaultdict(lambda: {
        'BULAN': '',
        'JENIS_PERKARA': '',
        'JUMLAH': 0,
        'PRA_PENUNTUTAN': 0,
        'PENUNTUTAN': 0,
        'UPAYA_HUKUM': 0
    })
    
    # Process data for charts by tahapan
    chart_data = {
        'pra_penuntutan': defaultdict(int),
        'penuntutan': defaultdict(int),
        'upaya_hukum': defaultdict(int)
    }
    
    # Normalisasi mapping untuk tahapan
    def normalize_tahapan(tahapan_text):
        if not tahapan_text:
            return 'pra_penuntutan'
        tahapan_upper = tahapan_text.upper().strip()
        if 'PRA' in tahapan_upper or 'PENYIDIKAN' in tahapan_upper:
            return 'pra_penuntutan'
        elif 'PENUNTUTAN' in tahapan_upper and 'PRA' not in tahapan_upper:
            return 'penuntutan'
        elif 'UPAYA' in tahapan_upper or 'HUKUM' in tahapan_upper or 'BANDING' in tahapan_upper or 'KASASI' in tahapan_upper:
            return 'upaya_hukum'
        else:
            return 'pra_penuntutan'
    
    # Normalisasi mapping untuk jenis perkara
    def normalize_jenis_perkara(jenis_text):
        if not jenis_text:
            return 'PERKARA LAINNYA'
        jenis_upper = jenis_text.upper().strip()
        if 'NARKOT' in jenis_upper or 'NARKOBA' in jenis_upper:
            return 'NARKOBA'
        elif 'ANAK' in jenis_upper:
            return 'PERKARA ANAK'
        elif 'SUSILA' in jenis_upper or 'KESUSILAAN' in jenis_upper:
            return 'KESUSILAAN'
        elif 'JUDI' in jenis_upper:
            return 'JUDI'
        elif 'KDRT' in jenis_upper or 'KEKERASAN DALAM RUMAH' in jenis_upper:
            return 'KDRT'
        elif 'OHARDA' in jenis_upper or 'HARDA' in jenis_upper:
            return 'OHARDA'
        else:
            return 'PERKARA LAINNYA'
    
    for row in rows:
        key = (row['bulan_nama'], row['jenis_perkara'])
        data_summary[key]['BULAN'] = row['bulan_nama']
        data_summary[key]['JENIS_PERKARA'] = row['jenis_perkara']
        
        # Normalisasi tahapan dan jenis perkara
        tahapan_normalized = normalize_tahapan(row['tahapan_penanganan'])
        jenis_normalized = normalize_jenis_perkara(row['jenis_perkara'])
        
        # Update data summary
        if tahapan_normalized == 'pra_penuntutan':
            data_summary[key]['PRA_PENUNTUTAN'] += 1
        elif tahapan_normalized == 'penuntutan':
            data_summary[key]['PENUNTUTAN'] += 1
        elif tahapan_normalized == 'upaya_hukum':
            data_summary[key]['UPAYA_HUKUM'] += 1
        
        # Update chart data dengan normalisasi
        chart_data[tahapan_normalized][jenis_normalized] += 1
        
        data_summary[key]['JUMLAH'] = (data_summary[key]['PRA_PENUNTUTAN'] +
                                     data_summary[key]['PENUNTUTAN'] +
                                     data_summary[key]['UPAYA_HUKUM'])
    
    # Ensure all predefined categories are present for each month
    for month in display_months:
        for category in predefined_categories:
            key = (month, category)
            if key not in data_summary:
                data_summary[key] = {
                    'BULAN': month,
                    'JENIS_PERKARA': category,
                    'JUMLAH': 0,
                    'PRA_PENUNTUTAN': 0,
                    'PENUNTUTAN': 0,
                    'UPAYA_HUKUM': 0
                }
    
    # Ensure all categories are present in chart data
    for category in predefined_categories:
        if category not in chart_data['pra_penuntutan']:
            chart_data['pra_penuntutan'][category] = 0
        if category not in chart_data['penuntutan']:
            chart_data['penuntutan'][category] = 0
        if category not in chart_data['upaya_hukum']:
            chart_data['upaya_hukum'][category] = 0
    
    # Convert to list and sort
    report_data = list(data_summary.values())
    
    # Sort by month order, then by category order
    month_order = {name: i for i, name in enumerate(month_names.values())}
    category_order = {cat: i for i, cat in enumerate(predefined_categories)}
    
    report_data.sort(key=lambda x: (month_order.get(x['BULAN'], 0), category_order.get(x['JENIS_PERKARA'], 0)))
    
    # Calculate totals
    total_keseluruhan = sum(item['JUMLAH'] for item in report_data)
    total_pra_penuntutan = sum(item['PRA_PENUNTUTAN'] for item in report_data)
    total_penuntutan = sum(item['PENUNTUTAN'] for item in report_data)
    total_upaya_hukum = sum(item['UPAYA_HUKUM'] for item in report_data)
    
    # Current date for display
    current_date = datetime.now().strftime("%d %B %Y")
    
    return render_template('laporan_pidum_bulanan.html',
                         report_data=report_data,
                         tahun=tahun,
                         bulan=bulan,
                         available_months=available_months,
                         chart_data=chart_data,
                         total_keseluruhan=total_keseluruhan,
                         total_pra_penuntutan=total_pra_penuntutan,
                         total_penuntutan=total_penuntutan,
                         total_upaya_hukum=total_upaya_hukum,
                         current_date=current_date)

@app.route('/input_pidum', methods=['GET', 'POST'])
@login_required
def input_pidum():
    # Redirect langsung ke halaman pilihan import data PIDUM
    # Bisa pilih dari dropdown import tahapan penanganan
    if request.method == 'GET':
        return render_template('pilihan_import_pidum.html')
    
    # Jika POST, proses input manual
    if request.method == 'POST':
        try:
            # Get form data
            data = {
                'NO': request.form['no'],
                'PERIODE': request.form['periode'],
                'TANGGAL': request.form['tanggal'],
                'JENIS PERKARA': request.form['jenis_perkara'],
                'TAHAPAN_PENANGANAN': 'PRA PENUNTUTAN',  # Default value for manual input
                'KETERANGAN': ''  # Default empty for manual input
            }
            
            # Insert to database
            insert_pidum_data(data)
            
            flash('Data PIDUM berhasil ditambahkan!', 'success')
            return redirect(url_for('input_pidum'))
        except Exception as e:
            flash(f'Terjadi kesalahan: {str(e)}', 'error')
    
    # Get data from database untuk fallback
    data = get_pidum_data_for_export()
    return render_template('pilihan_import_pidum.html', data=data)

@app.route('/manual_input_pidum', methods=['GET', 'POST'])
@login_required
def manual_input_pidum():
    """Route untuk input manual data PIDUM (form asli)"""
    if request.method == 'POST':
        try:
            # Get form data
            data = {
                'NO': request.form['no'],
                'PERIODE': request.form['periode'],
                'TANGGAL': request.form['tanggal'],
                'JENIS PERKARA': request.form['jenis_perkara'],
                'TAHAPAN_PENANGANAN': 'PRA PENUNTUTAN',  # Default value for manual input
                'KETERANGAN': request.form.get('keterangan', '')  # Get keterangan from form
            }
            
            # Insert to database
            insert_pidum_data(data)
            
            flash('Data PIDUM berhasil ditambahkan!', 'success')
            return redirect(url_for('manual_input_pidum'))
        except Exception as e:
            flash(f'Terjadi kesalahan: {str(e)}', 'error')
    
    # Get data from database
    data = get_pidum_data_for_export()
    return render_template('input_pidum.html', data=data)

@app.route('/input_pidsus', methods=['GET', 'POST'])
@login_required
def input_pidsus():
    if request.method == 'POST':
        try:
            # Get form data
            data = {
                'NO': request.form['no'],
                'PERIODE': request.form['periode'],
                'TANGGAL': request.form['tanggal'],
                'JENIS PERKARA': request.form['jenis_perkara'],
                'PENYIDIKAN': request.form['penyidikan'],
                'PENUNTUTAN': request.form['penuntutan'],
                'KETERANGAN': request.form['keterangan']
            }
            
            # Insert to database
            insert_pidsus_data(data)
            
            flash('Data PIDSUS berhasil ditambahkan!', 'success')
            return redirect(url_for('input_pidsus'))
        except Exception as e:
            flash(f'Terjadi kesalahan: {str(e)}', 'error')
    
    # Get data from database
    data = get_pidsus_data_for_export()
    return render_template('input_pidsus.html', data=data)

@app.route('/view_pidum')
@login_required
def view_pidum():
    data = get_all_pidum_data()  # Changed to include ID for delete functionality
    return render_template('view_pidum.html', data=data)

@app.route('/view_pidsus')
@login_required
def view_pidsus():
    data = get_pidsus_data_for_export()
    return render_template('view_pidsus.html', data=data)

@app.route('/delete_all_pidum', methods=['POST'])
@login_required
def delete_all_pidum():
    """Delete all PIDUM data"""
    try:
        deleted_count = delete_all_pidum_data()
        flash(f'Berhasil menghapus {deleted_count} data PIDUM', 'success')
    except Exception as e:
        flash(f'Error menghapus data: {str(e)}', 'error')
    return redirect(url_for('view_pidum'))

@app.route('/delete_pidum_item/<int:item_id>', methods=['POST'])
@login_required
def delete_pidum_item_route(item_id):
    """Delete single PIDUM item by ID"""
    try:
        success = delete_pidum_item(item_id)
        if success:
            flash('Data berhasil dihapus', 'success')
        else:
            flash('Data tidak ditemukan', 'error')
    except Exception as e:
        flash(f'Error menghapus data: {str(e)}', 'error')
    return redirect(url_for('view_pidum'))

@app.route('/laporan_pidum')
@login_required
def laporan_pidum():
    # Get filter parameters
    # Bulan bisa kosong untuk menampilkan semua bulan
    bulan_raw = request.args.get('bulan')
    bulan = int(bulan_raw) if bulan_raw not in (None, '') else None
    tahun = request.args.get('tahun', type=int)
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')

    # Default to current month/year if not specified
    from datetime import datetime
    # Hanya default ke bulan berjalan jika user tidak mengirim parameter 'bulan' sama sekali
    if 'bulan' not in request.args:
        if not bulan:
            bulan = datetime.now().month
    if not tahun:
        tahun = datetime.now().year

    # Query database using filters similar to laporan_pidum_new
    import sqlite3
    from collections import defaultdict

    conn = sqlite3.connect('db/kejaksaan.db')
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    where_conditions = []
    params = []

    if bulan:
        where_conditions.append("strftime('%m', tanggal) = ?")
        params.append(f"{bulan:02d}")

    if tahun:
        where_conditions.append("strftime('%Y', tanggal) = ?")
        params.append(str(tahun))

    if start_date:
        where_conditions.append("tanggal >= ?")
        params.append(start_date)

    if end_date:
        where_conditions.append("tanggal <= ?")
        params.append(end_date)

    where_clause = ""
    if where_conditions:
        where_clause = "WHERE " + " AND ".join(where_conditions)

    query = f"""
    SELECT jenis_perkara, tahapan_penanganan,
           CASE strftime('%m', tanggal)
               WHEN '01' THEN 'Januari'
               WHEN '02' THEN 'Februari'
               WHEN '03' THEN 'Maret'
               WHEN '04' THEN 'April'
               WHEN '05' THEN 'Mei'
               WHEN '06' THEN 'Juni'
               WHEN '07' THEN 'Juli'
               WHEN '08' THEN 'Agustus'
               WHEN '09' THEN 'September'
               WHEN '10' THEN 'Oktober'
               WHEN '11' THEN 'November'
               WHEN '12' THEN 'Desember'
           END as bulan_nama
    FROM pidum_data {where_clause}
    ORDER BY jenis_perkara
    """

    cursor.execute(query, params)
    rows = cursor.fetchall()
    conn.close()

    # Aggregate per jenis_perkara
    agg = defaultdict(lambda: {
        'jenis_perkara': '',
        'BULAN': '',
        'JUMLAH': 0,
        'PRA PENUNTUTAN': 0,
        'PENUNTUTAN': 0,
        'UPAYA HUKUM': 0
    })

    for row in rows:
        jenis = (row['jenis_perkara'] or '').strip()
        tahapan = (row['tahapan_penanganan'] or '').upper()
        bulan_nama_row = row['bulan_nama']

        if not jenis:
            jenis = 'Tidak Diketahui'
        
        # Group by both month and jenis_perkara
        agg_key = (bulan_nama_row, jenis)

        agg[agg_key]['jenis_perkara'] = jenis
        agg[agg_key]['BULAN'] = bulan_nama_row

        if 'PRA PENUNTUTAN' in tahapan:
            agg[agg_key]['PRA PENUNTUTAN'] += 1
        elif 'PENUNTUTAN' in tahapan:
            agg[agg_key]['PENUNTUTAN'] += 1
        elif 'UPAYA HUKUM' in tahapan:
            agg[agg_key]['UPAYA HUKUM'] += 1
        else:
            agg[agg_key]['PRA PENUNTUTAN'] += 1
        
        agg[agg_key]['JUMLAH'] = (
            agg[agg_key]['PRA PENUNTUTAN'] + agg[agg_key]['PENUNTUTAN'] + agg[agg_key]['UPAYA HUKUM']
        )

    # Ensure all predefined categories exist (including zeros)
    predefined_categories = [
        'NARKOBA',
        'PERKARA ANAK',
        'KESUSILAAN',
        'JUDI',
        'KDRT',
        'OHARDA',
        'PERKARA LAINNYA'
    ]

    # Normalize keys to predefined mapping
    def map_to_predefined(name: str) -> str:
        upper = (name or '').upper()
        if 'NARKOBA' in upper:
            return 'NARKOBA'
        if 'ANAK' in upper:
            return 'PERKARA ANAK'
        if 'KESUSILAAN' in upper or 'SUSILA' in upper:
            return 'KESUSILAAN'
        if 'JUDI' in upper:
            return 'JUDI'
        if 'KDRT' in upper:
            return 'KDRT'
        if 'OHARDA' in upper:
            return 'OHARDA'
        return 'PERKARA LAINNYA'

    # Re-map aggregated keys into predefined buckets
    remapped = defaultdict(lambda: {
        'jenis_perkara': '',
        'BULAN': '',
        'JUMLAH': 0,
        'PRA PENUNTUTAN': 0,
        'PENUNTUTAN': 0,
        'UPAYA HUKUM': 0
    })

    for (bulan_nama_row, jenis), data in agg.items():
        bucket = map_to_predefined(jenis)
        new_key = (bulan_nama_row, bucket)

        remapped[new_key]['BULAN'] = bulan_nama_row
        remapped[new_key]['jenis_perkara'] = bucket
        remapped[new_key]['PRA PENUNTUTAN'] += data['PRA PENUNTUTAN']
        remapped[new_key]['PENUNTUTAN'] += data['PENUNTUTAN']
        remapped[new_key]['UPAYA HUKUM'] += data['UPAYA HUKUM']
        remapped[new_key]['JUMLAH'] += data['JUMLAH']

    # Ensure all predefined categories are present for each month
    # Get all months present in the data
    all_months = set()
    for key in remapped.keys():
        all_months.add(key[0])  # key[0] is the month name
    
    # If no specific month is selected, use all months
    if not bulan:
        month_names = {
            1: 'Januari', 2: 'Februari', 3: 'Maret', 4: 'April',
            5: 'Mei', 6: 'Juni', 7: 'Juli', 8: 'Agustus',
            9: 'September', 10: 'Oktober', 11: 'November', 12: 'Desember'
        }
        all_months = {month_names.get(bulan, 'Januari') for bulan in range(1, 13)}
    else:
        month_names = {
            1: 'Januari', 2: 'Februari', 3: 'Maret', 4: 'April',
            5: 'Mei', 6: 'Juni', 7: 'Juli', 8: 'Agustus',
            9: 'September', 10: 'Oktober', 11: 'November', 12: 'Desember'
        }
        all_months = {month_names.get(bulan, 'Januari')}
    
    # Add missing categories with zero values for each month
    for month in all_months:
        for category in predefined_categories:
            key = (month, category)
            if key not in remapped:
                remapped[key] = {
                    'BULAN': month,
                    'jenis_perkara': category,
                    'JUMLAH': 0,
                    'PRA PENUNTUTAN': 0,
                    'PENUNTUTAN': 0,
                    'UPAYA HUKUM': 0
                }

    # Convert to list with sequential NO, ordered by predefined list
    report_data = list(remapped.values())

    month_map = {name: num for num, name in month_names.items()}
    
    # Sort data
    # If a specific month is selected, sort by category only
    if bulan:
        report_data.sort(key=lambda x: predefined_categories.index(x['jenis_perkara']) if x['jenis_perkara'] in predefined_categories else -1)
    else:
        # Sort by month, then by category
        report_data.sort(key=lambda x: (month_map.get(x['BULAN'], 0), predefined_categories.index(x['jenis_perkara']) if x['jenis_perkara'] in predefined_categories else -1))

    # Add sequential NO
    for i, item in enumerate(report_data):
        item['NO'] = i + 1

    # Calculate totals
    total_pra_penuntutan = sum(item['PRA PENUNTUTAN'] for item in report_data)
    total_penuntutan = sum(item['PENUNTUTAN'] for item in report_data)
    total_upaya_hukum = sum(item['UPAYA HUKUM'] for item in report_data)
    total_keseluruhan = sum(item['JUMLAH'] for item in report_data)

    return render_template(
        'laporan_pidum.html',
        report_data=report_data,
        bulan=bulan,
        tahun=tahun,
        start_date=start_date,
        end_date=end_date,
        bulan_nama=(month_names.get(bulan) if bulan else 'Semua Bulan'),
        total_pra_penuntutan=total_pra_penuntutan,
        total_penuntutan=total_penuntutan,
        total_upaya_hukum=total_upaya_hukum,
        total_keseluruhan=total_keseluruhan
    )

@app.route('/laporan_pidum_new')
@login_required
def laporan_pidum_new():
    # Get filter parameters
    bulan = request.args.get('bulan', type=int)
    tahun = request.args.get('tahun', type=int, default=2025)
    periode_filter = request.args.get('periode')
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    
    from datetime import datetime
    import sqlite3
    from collections import defaultdict
    
    # Get database connection
    conn = sqlite3.connect('db/kejaksaan.db')
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    # Build query based on filters
    where_conditions = []
    params = []
    
    if bulan:
        where_conditions.append("strftime('%m', tanggal) = ?")
        params.append(f"{bulan:02d}")
    
    if tahun:
        where_conditions.append("strftime('%Y', tanggal) = ?")
        params.append(str(tahun))
    
    if periode_filter:
        where_conditions.append("periode = ?")
        params.append(periode_filter)
    
    if start_date:
        where_conditions.append("tanggal >= ?")
        params.append(start_date)
    
    if end_date:
        where_conditions.append("tanggal <= ?")
        params.append(end_date)
    
    where_clause = ""
    if where_conditions:
        where_clause = "WHERE " + " AND ".join(where_conditions)
    
    # Get data from pidum_data table
    query = f"""
    SELECT periode, jenis_perkara, tanggal, tahapan_penanganan
    FROM pidum_data {where_clause}
    ORDER BY tanggal DESC, periode, jenis_perkara
    """
    
    cursor.execute(query, params)
    rows = cursor.fetchall()
    
    # Get unique periods for filter dropdown
    cursor.execute("SELECT DISTINCT periode FROM pidum_data ORDER BY periode")
    periode_options = [row[0] for row in cursor.fetchall() if row[0]]
    
    conn.close()
    
    # Process data for report
    data_summary = defaultdict(lambda: {
        'PERIODE': '',
        'JENIS_PERKARA': '',
        'TANGGAL': '',
        'TOTAL': 0,
        'PRA_PENUNTUTAN': 0,
        'PENUNTUTAN': 0,
        'UPAYA_HUKUM': 0
    })
    
    for row in rows:
        key = (row['periode'], row['jenis_perkara'], row['tanggal'])
        data_summary[key]['PERIODE'] = row['periode']
        data_summary[key]['JENIS_PERKARA'] = row['jenis_perkara']
        data_summary[key]['TANGGAL'] = row['tanggal']
        
        # Map tahapan_penanganan to report columns
        tahapan = row['tahapan_penanganan'].upper() if row['tahapan_penanganan'] else ''
        
        if 'PRA PENUNTUTAN' in tahapan:
            data_summary[key]['PRA_PENUNTUTAN'] += 1
        elif 'PENUNTUTAN' in tahapan:
            data_summary[key]['PENUNTUTAN'] += 1
        elif 'UPAYA HUKUM' in tahapan:
            data_summary[key]['UPAYA_HUKUM'] += 1
        else:
            # Default ke pra penuntutan jika tidak jelas
            data_summary[key]['PRA_PENUNTUTAN'] += 1
        
        data_summary[key]['TOTAL'] = (data_summary[key]['PRA_PENUNTUTAN'] + 
                                     data_summary[key]['PENUNTUTAN'] + 
                                     data_summary[key]['UPAYA_HUKUM'])
    
    # Convert to list
    report_data = list(data_summary.values())
    
    # Calculate totals
    total_keseluruhan = sum(item['TOTAL'] for item in report_data)
    total_pra_penuntutan = sum(item['PRA_PENUNTUTAN'] for item in report_data)
    total_penuntutan = sum(item['PENUNTUTAN'] for item in report_data)
    total_upaya_hukum = sum(item['UPAYA_HUKUM'] for item in report_data)
    
    # Current date for display
    current_date = datetime.now().strftime("%d %B %Y")
    
    return render_template('laporan_pidum_new.html',
                         report_data=report_data,
                         bulan=bulan,
                         tahun=tahun,
                         periode_filter=periode_filter,
                         periode_options=periode_options,
                         start_date=start_date,
                         end_date=end_date,
                         total_keseluruhan=total_keseluruhan,
                         total_pra_penuntutan=total_pra_penuntutan,
                         total_penuntutan=total_penuntutan,
                         total_upaya_hukum=total_upaya_hukum,
                         current_date=current_date)

@app.route('/export_pidum_excel')
@login_required
def export_pidum_excel():
    data = get_pidum_data_for_export()
    if not data:
        flash('Tidak ada data PIDUM untuk diekspor', 'warning')
        return redirect(url_for('view_pidum'))
    
    # Create DataFrame
    df = pd.DataFrame(data)
    
    # Create Excel file
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='Data PIDUM', index=False)
        
        # Get the workbook and worksheet
        workbook = writer.book
        worksheet = writer.sheets['Data PIDUM']
        
        # Style the header
        header_font = Font(bold=True, color='FFFFFF')
        header_fill = PatternFill(start_color='4F81BD', end_color='4F81BD', fill_type='solid')
        header_alignment = Alignment(horizontal='center', vertical='center')
        
        for cell in worksheet[1]:  # First row
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
        
        # Auto-adjust column width
        for column in worksheet.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = (max_length + 2)
            worksheet.column_dimensions[column_letter].width = adjusted_width
    
    output.seek(0)
    
    # Generate filename with timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"data_pidum_{timestamp}.xlsx"
    
    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )

@app.route('/export_pidum_new_excel')
@login_required
def export_pidum_new_excel():
    # Get filter parameters from request
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    bulan = request.args.get('bulan', type=int)
    tahun = request.args.get('tahun', type=int, default=2025)
    periode_filter = request.args.get('periode')
    jenis_perkara_filter = request.args.get('jenis_perkara')
    tahapan_filter = request.args.get('tahapan')
    
    import sqlite3
    from collections import defaultdict
    
    # Get database connection
    conn = sqlite3.connect('db/kejaksaan.db')
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    # Build WHERE clause based on filters
    where_conditions = []
    params = []
    
    if start_date and end_date:
        where_conditions.append("DATE(tanggal) BETWEEN ? AND ?")
        params.extend([start_date, end_date])
    elif bulan and tahun:
        where_conditions.append("strftime('%m', tanggal) = ? AND strftime('%Y', tanggal) = ?")
        params.extend([f"{bulan:02d}", str(tahun)])
    elif tahun:
        where_conditions.append("strftime('%Y', tanggal) = ?")
        params.append(str(tahun))
    
    if periode_filter:
        where_conditions.append("periode = ?")
        params.append(periode_filter)
    
    if jenis_perkara_filter:
        where_conditions.append("jenis_perkara = ?")
        params.append(jenis_perkara_filter)
    
    if tahapan_filter:
        where_conditions.append("tahapan_penanganan = ?")
        params.append(tahapan_filter)
    
    where_clause = " AND ".join(where_conditions) if where_conditions else "1=1"
    
    # Get filtered data
    query = f"""
    SELECT no, periode, tanggal, jenis_perkara, tahapan_penanganan, keterangan, created_at
    FROM pidum_data 
    WHERE {where_clause}
    ORDER BY tanggal DESC, id DESC
    """
    
    cursor.execute(query, params)
    rows = cursor.fetchall()
    
    # Convert to list of dictionaries
    data = []
    for row in rows:
        data.append({
            'No': row['no'],
            'Periode': row['periode'],
            'Tanggal Transaksi': row['tanggal'],
            'Jenis Perkara': row['jenis_perkara'],
            'Tahapan Penanganan': row['tahapan_penanganan'],
            'Keterangan': row['keterangan'],
            'Tanggal Input': row['created_at']
        })
    
    conn.close()
    
    if not data:
        flash('Tidak ada data untuk filter yang dipilih', 'warning')
        return redirect(url_for('laporan_pidum_new'))
    
    # Create DataFrame
    df = pd.DataFrame(data)
    
    # Create Excel file with chart
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='Data PIDUM', index=False, startrow=6)
        
        # Get the workbook and worksheet
        workbook = writer.book
        worksheet = writer.sheets['Data PIDUM']
        
        # Add title and filter info
        worksheet['A1'] = 'LAPORAN PIDUM PER TANGGAL'
        worksheet['A1'].font = Font(size=16, bold=True)
        worksheet['A2'] = f'Tahun: {tahun}'
        if bulan:
            bulan_nama = ['', 'Januari', 'Februari', 'Maret', 'April', 'Mei', 'Juni', 
                         'Juli', 'Agustus', 'September', 'Oktober', 'November', 'Desember'][bulan]
            worksheet['A2'] = f'Tahun: {tahun}, Bulan: {bulan_nama}'
        if start_date and end_date:
            worksheet['A3'] = f'Periode: {start_date} s/d {end_date}'
        worksheet['A4'] = f'Total Data: {len(data)} record'
        worksheet['A5'] = f'Dicetak pada: {datetime.now().strftime("%d %B %Y %H:%M")}'
        
        # Style the header (row 7)
        header_font = Font(bold=True, color='FFFFFF')
        header_fill = PatternFill(start_color='4F81BD', end_color='4F81BD', fill_type='solid')
        header_alignment = Alignment(horizontal='center', vertical='center')
        
        for cell in worksheet[7]:  # Header row
            if cell.value:  # Only style cells with content
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment
        
        # Auto-adjust column width
        for column in worksheet.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min((max_length + 2), 50)  # Max width 50
            worksheet.column_dimensions[column_letter].width = adjusted_width
    
    output.seek(0)
    
    # Generate filename with filter info
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filter_info = f"tahun_{tahun}"
    if bulan:
        filter_info += f"_bulan_{bulan}"
    if start_date and end_date:
        filter_info += f"_{start_date}_to_{end_date}"
    
    filename = f"laporan_pidum_per_tanggal_{filter_info}_{timestamp}.xlsx"
    
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=filename
    )

@app.route('/export_pidum_new_word')
@login_required
def export_pidum_new_word():
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        import matplotlib.pyplot as plt
        import matplotlib
        matplotlib.use('Agg')  # Use non-interactive backend
        import numpy as np
        from collections import defaultdict
    except ImportError as e:
        flash(f'Library tidak tersedia: {str(e)}. Install dengan: pip install python-docx matplotlib', 'error')
        return redirect(url_for('laporan_pidum_new'))
    
    # Get filter parameters from request
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    bulan = request.args.get('bulan', type=int)
    tahun = request.args.get('tahun', type=int, default=2025)
    periode_filter = request.args.get('periode')
    jenis_perkara_filter = request.args.get('jenis_perkara')
    tahapan_filter = request.args.get('tahapan')
    
    import sqlite3
    
    # Get database connection
    conn = sqlite3.connect('db/kejaksaan.db')
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    # Build WHERE clause based on filters
    where_conditions = []
    params = []
    
    if start_date and end_date:
        where_conditions.append("DATE(tanggal) BETWEEN ? AND ?")
        params.extend([start_date, end_date])
    elif bulan and tahun:
        where_conditions.append("strftime('%m', tanggal) = ? AND strftime('%Y', tanggal) = ?")
        params.extend([f"{bulan:02d}", str(tahun)])
    elif tahun:
        where_conditions.append("strftime('%Y', tanggal) = ?")
        params.append(str(tahun))
    
    if periode_filter:
        where_conditions.append("periode = ?")
        params.append(periode_filter)
    
    if jenis_perkara_filter:
        where_conditions.append("jenis_perkara = ?")
        params.append(jenis_perkara_filter)
    
    if tahapan_filter:
        where_conditions.append("tahapan_penanganan = ?")
        params.append(tahapan_filter)
    
    where_clause = " AND ".join(where_conditions) if where_conditions else "1=1"
    
    # Get filtered data
    query = f"""
    SELECT no, periode, tanggal, jenis_perkara, tahapan_penanganan, keterangan, created_at
    FROM pidum_data 
    WHERE {where_clause}
    ORDER BY tanggal DESC, id DESC
    """
    
    cursor.execute(query, params)
    rows = cursor.fetchall()
    conn.close()
    
    if not rows:
        flash('Tidak ada data untuk filter yang dipilih', 'warning')
        return redirect(url_for('laporan_pidum_new'))
    
    # Process data for chart
    chart_data = defaultdict(int)
    all_jenis_perkara = ['NARKOBA', 'PERKARA ANAK', 'KESUSILAAN', 'JUDI', 'KDRT', 'OHARDA', 'PERKARA LAINNYA']
    
    # Normalisasi jenis perkara
    def normalize_jenis_perkara(jenis_text):
        if not jenis_text:
            return 'PERKARA LAINNYA'
        jenis_upper = jenis_text.upper().strip()
        if 'NARKOT' in jenis_upper or 'NARKOBA' in jenis_upper:
            return 'NARKOBA'
        elif 'ANAK' in jenis_upper:
            return 'PERKARA ANAK'
        elif 'SUSILA' in jenis_upper or 'KESUSILAAN' in jenis_upper:
            return 'KESUSILAAN'
        elif 'JUDI' in jenis_upper:
            return 'JUDI'
        elif 'KDRT' in jenis_upper or 'KEKERASAN DALAM RUMAH' in jenis_upper:
            return 'KDRT'
        elif 'OHARDA' in jenis_upper or 'HARDA' in jenis_upper:
            return 'OHARDA'
        else:
            return 'PERKARA LAINNYA'
    
    # Count data for chart
    for row in rows:
        normalized_jenis = normalize_jenis_perkara(row['jenis_perkara'])
        chart_data[normalized_jenis] += 1
    
    # Ensure all categories are present
    chart_values = [chart_data.get(jenis, 0) for jenis in all_jenis_perkara]
    
    # Create chart using matplotlib
    plt.figure(figsize=(12, 8))
    colors = ['#dc3545', '#28a745', '#ffc107', '#17a2b8', '#6f42c1', '#fd7e14', '#6c757d']
    
    bars = plt.bar(all_jenis_perkara, chart_values, color=colors)
    plt.title('Distribusi Kasus per Jenis Perkara (Berdasarkan Filter Aktif)', fontsize=16, fontweight='bold', pad=20)
    plt.xlabel('Jenis Perkara', fontsize=12, fontweight='bold')
    plt.ylabel('Jumlah Kasus', fontsize=12, fontweight='bold')
    plt.xticks(rotation=45, ha='right')
    plt.grid(axis='y', alpha=0.3)
    
    # Add value labels on bars
    for bar, value in zip(bars, chart_values):
        if value > 0:
            plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.1, 
                    str(value), ha='center', va='bottom', fontweight='bold')
    
    plt.tight_layout()
    
    # Save chart to temporary file
    import tempfile
    chart_temp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
    plt.savefig(chart_temp.name, dpi=300, bbox_inches='tight')
    plt.close()
    
    # Create Word document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add header
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_p.add_run('LAPORAN PIDUM PER TANGGAL')
    header_run.font.size = Pt(16)
    header_run.font.bold = True
    
    subheader_p = doc.add_paragraph()
    subheader_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subheader_run = subheader_p.add_run('KEJAKSAAN NEGERI')
    subheader_run.font.size = Pt(12)
    
    doc.add_paragraph()  # Empty line
    
    # Add filter information
    info_p = doc.add_paragraph()
    info_text = f"Tahun: {tahun}"
    if bulan:
        bulan_nama = ['', 'Januari', 'Februari', 'Maret', 'April', 'Mei', 'Juni', 
                     'Juli', 'Agustus', 'September', 'Oktober', 'November', 'Desember'][bulan]
        info_text = f"Tahun: {tahun}, Bulan: {bulan_nama}"
    if start_date and end_date:
        info_text += f", Periode: {start_date} s/d {end_date}"
    if periode_filter:
        info_text += f", Periode: {periode_filter}"
    if jenis_perkara_filter:
        info_text += f", Jenis Perkara: {jenis_perkara_filter}"
    if tahapan_filter:
        info_text += f", Tahapan: {tahapan_filter}"
    
    info_run = info_p.add_run(f"Informasi Filter: {info_text}")
    info_run.font.bold = True
    
    total_p = doc.add_paragraph()
    total_run = total_p.add_run(f"Total Data: {len(rows)} record")
    total_run.font.bold = True
    
    date_p = doc.add_paragraph()
    date_run = date_p.add_run(f"Dicetak pada: {datetime.now().strftime('%d %B %Y %H:%M')}")
    date_run.font.bold = True
    
    doc.add_paragraph()  # Empty line
    
    # Add chart
    chart_p = doc.add_paragraph()
    chart_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    chart_title = chart_p.add_run('Grafik Distribusi Jenis Perkara')
    chart_title.font.size = Pt(14)
    chart_title.font.bold = True
    
    chart_para = doc.add_paragraph()
    chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(chart_temp.name, width=Inches(6))
    
    doc.add_paragraph()  # Empty line
    
    # Add table
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set header row
    header_cells = table.rows[0].cells
    headers = ['No', 'Periode', 'Tanggal', 'Jenis Perkara', 'Tahapan Penanganan', 'Keterangan', 'Tanggal Input']
    
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Make header bold
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add data rows
    for row in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['no'])
        row_cells[1].text = str(row['periode'])
        row_cells[2].text = str(row['tanggal'])
        row_cells[3].text = str(row['jenis_perkara'])
        row_cells[4].text = str(row['tahapan_penanganan'])
        row_cells[5].text = str(row['keterangan'])
        row_cells[6].text = str(row['created_at'])
        
        # Center align some columns
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Set column widths
    widths = [Inches(0.5), Inches(1), Inches(1), Inches(1.5), Inches(1.5), Inches(2), Inches(1.5)]
    for i, width in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = width
    
    # Add total row
    total_row = table.add_row().cells
    total_row[0].text = "TOTAL KESELURUHAN"
    total_row[1].text = str(len(rows))
    
    # Make total row bold and merge cells
    for i in range(2):
        for paragraph in total_row[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        total_row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Merge remaining cells in total row
    for i in range(2, 7):
        total_row[i].text = ""
    
    # Clean up temporary chart file
    import os
    os.unlink(chart_temp.name)
    
    # Save to BytesIO
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    # Generate filename with filter info
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filter_info = f"tahun_{tahun}"
    if bulan:
        filter_info += f"_bulan_{bulan}"
    if start_date and end_date:
        filter_info += f"_{start_date}_to_{end_date}"
    
    filename = f"laporan_pidum_per_tanggal_{filter_info}_{timestamp}.docx"
    
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=filename
    )

@app.route('/export_pidsus_excel')
@login_required
def export_pidsus_excel():
    data = get_pidsus_data_for_export()
    if not data:
        flash('Tidak ada data PIDSUS untuk diekspor', 'warning')
        return redirect(url_for('view_pidsus'))
    
    # Create DataFrame
    df = pd.DataFrame(data)
    
    # Create Excel file
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='Data PIDSUS', index=False)
        
        # Get the workbook and worksheet
        workbook = writer.book
        worksheet = writer.sheets['Data PIDSUS']
        
        # Style the header
        header_font = Font(bold=True, color='FFFFFF')
        header_fill = PatternFill(start_color='4F81BD', end_color='4F81BD', fill_type='solid')
        header_alignment = Alignment(horizontal='center', vertical='center')
        
        for cell in worksheet[1]:  # First row
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
        
        # Auto-adjust column width
        for column in worksheet.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = (max_length + 2)
            worksheet.column_dimensions[column_letter].width = adjusted_width
    
    output.seek(0)
    
    # Generate filename with timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"data_pidsus_{timestamp}.xlsx"
    
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=filename
    )

@app.route('/pidum_charts')
@login_required
def pidum_charts():
    data = get_pidum_data_for_export()
    if not data:
        flash('Tidak ada data PIDUM untuk ditampilkan dalam grafik', 'warning')
        return redirect(url_for('view_pidum'))
    
    # Create charts
    charts = {}
    
    # Chart by Jenis Perkara
    df = pd.DataFrame(data)
    jenis_perkara_counts = df['JENIS PERKARA'].value_counts()
    
    plt.figure(figsize=(10, 6))
    jenis_perkara_counts.plot(kind='bar', color='skyblue')
    plt.title('Jumlah Kasus Berdasarkan Jenis Perkara (PIDUM)')
    plt.xlabel('Jenis Perkara')
    plt.ylabel('Jumlah Kasus')
    plt.xticks(rotation=45)
    plt.tight_layout()
    
    img = io.BytesIO()
    plt.savefig(img, format='png')
    img.seek(0)
    charts['jenis_perkara'] = base64.b64encode(img.getvalue()).decode('utf-8')
    plt.close()
    
    # Chart by Upaya Hukum
    upaya_hukum_counts = df['UPAYA HUKUM'].value_counts()
    
    plt.figure(figsize=(8, 6))
    upaya_hukum_counts.plot(kind='pie', autopct='%1.1f%%', colors=['lightcoral', 'lightgreen', 'lightblue', 'lightyellow'])
    plt.title('Distribusi Upaya Hukum (PIDUM)')
    plt.tight_layout()
    
    img = io.BytesIO()
    plt.savefig(img, format='png')
    img.seek(0)
    charts['upaya_hukum'] = base64.b64encode(img.getvalue()).decode('utf-8')
    plt.close()
    
    return render_template('pidum_charts.html', charts=charts)

@app.route('/pidsus_charts')
@login_required
def pidsus_charts():
    data = get_pidsus_data_for_export()
    if not data:
        flash('Tidak ada data PIDSUS untuk ditampilkan dalam grafik', 'warning')
        return redirect(url_for('view_pidsus'))
    
    # Create charts
    charts = {}
    
    # Chart by Jenis Perkara
    df = pd.DataFrame(data)
    jenis_perkara_counts = df['JENIS PERKARA'].value_counts()
    
    plt.figure(figsize=(10, 6))
    jenis_perkara_counts.plot(kind='bar', color='skyblue')
    plt.title('Jumlah Kasus Berdasarkan Jenis Perkara (PIDSUS)')
    plt.xlabel('Jenis Perkara')
    plt.ylabel('Jumlah Kasus')
    plt.xticks(rotation=45)
    plt.tight_layout()
    
    img = io.BytesIO()
    plt.savefig(img, format='png')
    img.seek(0)
    charts['jenis_perkara'] = base64.b64encode(img.getvalue()).decode('utf-8')
    plt.close()
    
    # Chart by Penyidikan
    penyidikan_counts = df['PENYIDIKAN'].value_counts()
    
    plt.figure(figsize=(8, 6))
    penyidikan_counts.plot(kind='pie', autopct='%1.1f%%', colors=['lightcoral', 'lightgreen', 'lightblue', 'lightyellow'])
    plt.title('Distribusi Penyidikan (PIDSUS)')
    plt.tight_layout()
    
    img = io.BytesIO()
    plt.savefig(img, format='png')
    img.seek(0)
    charts['penyidikan'] = base64.b64encode(img.getvalue()).decode('utf-8')
    plt.close()
    
    return render_template('pidsus_charts.html', charts=charts)

@app.route('/database_info')
@login_required
def database_info():
    """Route untuk melihat informasi database"""
    stats = get_database_stats()
    return f"""
    <h2>Database Information</h2>
    <p><strong>Database Path:</strong> {stats['database_path']}</p>
    <p><strong>PIDUM Records:</strong> {stats['pidum_count']}</p>
    <p><strong>PIDSUS Records:</strong> {stats['pidsus_count']}</p>
    <p><strong>Total Records:</strong> {stats['pidum_count'] + stats['pidsus_count']}</p>
    <p><a href="/">Back to Home</a></p>
    """

@app.route('/import_pra_penuntutan_api', methods=['GET', 'POST'])
@login_required
def import_pra_penuntutan_api():
    """API khusus untuk import data pra penuntutan dengan format CSV khusus"""
    if request.method == 'POST':
        # Check if file is uploaded
        if 'file' not in request.files:
            flash('Tidak ada file yang dipilih', 'error')
            return redirect(url_for('import_tahapan', tahapan='pra_penuntutan'))
        
        file = request.files['file']
        if file.filename == '':
            flash('Tidak ada file yang dipilih', 'error')
            return redirect(url_for('import_tahapan', tahapan='pra_penuntutan'))
        
        if file and allowed_file_pra_penuntutan(file.filename):
            # Process the uploaded file with pra penuntutan specific handler
            result = process_pra_penuntutan_import_file(file)
            
            if result['success']:
                # Store import data in session for preview
                session['import_data_pra_penuntutan'] = result['data']
                session['import_filename_pra_penuntutan'] = file.filename
                session['import_format_type'] = 'pra_penuntutan'
                
                return render_template('import_pra_penuntutan_preview.html', 
                                     import_data=result['data'],
                                     filename=file.filename,
                                     total_rows=result['total_rows'],
                                     tahapan_penanganan='PRA PENUNTUTAN')
            else:
                flash(f'Error memproses file: {result["error"]}', 'error')
                return redirect(url_for('import_tahapan', tahapan='pra_penuntutan'))
        else:
            flash('Format file tidak didukung. Gunakan CSV, XLS, atau XLSX.', 'error')
            return redirect(url_for('import_tahapan', tahapan='pra_penuntutan'))
    
    # GET request - show upload form
    return render_template('import_pra_penuntutan.html')

@app.route('/confirm_import_pra_penuntutan', methods=['POST'])
@login_required
def confirm_import_pra_penuntutan():
    """Confirm and process pra penuntutan import"""
    import_data = session.get('import_data_pra_penuntutan', [])
    
    if not import_data:
        flash('Tidak ada data import yang tersedia', 'error')
        return redirect(url_for('import_tahapan', tahapan='pra_penuntutan'))
    
    try:
        # Prepare data from form
        prepared_data = prepare_pra_penuntutan_data_for_db(import_data, request.form)
        
        # Insert data to database
        success_count = 0
        error_count = 0
        error_details = []
        
        for data in prepared_data:
            try:
                insert_pidum_data(data)
                success_count += 1
            except Exception as e:
                error_count += 1
                error_details.append(f"Baris {data.get('NO', '?')}: {str(e)}")
                print(f"Error inserting row {data.get('NO', '?')}: {e}")
        
        # Clear session data
        session.pop('import_data_pra_penuntutan', None)
        session.pop('import_filename_pra_penuntutan', None)
        session.pop('import_format_type', None)
        
        # Show results
        if success_count > 0:
            flash(f'Berhasil import {success_count} data PRA PENUNTUTAN', 'success')
            if error_count > 0:
                flash(f'{error_count} data gagal diimport', 'warning')
                # Show some error details
                for error in error_details[:3]:
                    flash(error, 'warning')
        else:
            flash('Tidak ada data yang berhasil diimport', 'error')
            if error_details:
                for error in error_details[:3]:
                    flash(error, 'error')
            
    except Exception as e:
        flash(f'Error saat import data: {str(e)}', 'error')
    
    return redirect(url_for('view_pidum'))

@app.route('/import_upaya_hukum_api', methods=['GET', 'POST'])
@login_required
def import_upaya_hukum_api():
    """API khusus untuk import data upaya hukum dengan format CSV khusus"""
    if request.method == 'POST':
        # Check if file is uploaded
        if 'file' not in request.files:
            flash('Tidak ada file yang dipilih', 'error')
            return redirect(url_for('import_tahapan', tahapan='upaya_hukum'))
        
        file = request.files['file']
        if file.filename == '':
            flash('Tidak ada file yang dipilih', 'error')
            return redirect(url_for('import_tahapan', tahapan='upaya_hukum'))
        
        if file and allowed_file_upaya_hukum(file.filename):
            # Process the uploaded file with upaya hukum specific handler
            result = process_upaya_hukum_import_file(file)
            
            if result['success']:
                # Store import data in session for preview
                session['import_data_upaya_hukum'] = result['data']
                session['import_filename_upaya_hukum'] = file.filename
                session['import_format_type'] = 'upaya_hukum'
                
                return render_template('import_upaya_hukum_preview.html', 
                                     import_data=result['data'],
                                     filename=file.filename,
                                     total_rows=result['total_rows'],
                                     tahapan_penanganan='UPAYA HUKUM')
            else:
                flash(f'Error memproses file: {result["error"]}', 'error')
                return redirect(url_for('import_tahapan', tahapan='upaya_hukum'))
        else:
            flash('Format file tidak didukung. Gunakan CSV, XLS, atau XLSX.', 'error')
            return redirect(url_for('import_tahapan', tahapan='upaya_hukum'))
    
    # GET request - show upload form
    return render_template('import_upaya_hukum.html')

@app.route('/confirm_import_upaya_hukum', methods=['POST'])
@login_required
def confirm_import_upaya_hukum():
    """Confirm and process upaya hukum import"""
    import_data = session.get('import_data_upaya_hukum', [])
    
    if not import_data:
        flash('Tidak ada data import yang tersedia', 'error')
        return redirect(url_for('import_tahapan', tahapan='upaya_hukum'))
    
    try:
        # Prepare data from form
        prepared_data = prepare_upaya_hukum_data_for_db(import_data, request.form)
        
        # Insert data to database
        success_count = 0
        error_count = 0
        error_details = []
        
        for data in prepared_data:
            try:
                insert_pidum_data(data)
                success_count += 1
            except Exception as e:
                error_count += 1
                error_details.append(f"Baris {data.get('NO', '?')}: {str(e)}")
                print(f"Error inserting row {data.get('NO', '?')}: {e}")
        
        # Clear session data
        session.pop('import_data_upaya_hukum', None)
        session.pop('import_filename_upaya_hukum', None)
        session.pop('import_format_type', None)
        
        # Show results
        if success_count > 0:
            flash(f'Berhasil import {success_count} data UPAYA HUKUM', 'success')
            if error_count > 0:
                flash(f'{error_count} data gagal diimport', 'warning')
                # Show some error details
                for error in error_details[:3]:
                    flash(error, 'warning')
        else:
            flash('Tidak ada data yang berhasil diimport', 'error')
            if error_details:
                for error in error_details[:3]:
                    flash(error, 'error')
            
    except Exception as e:
        flash(f'Error saat import data: {str(e)}', 'error')
    
    return redirect(url_for('view_pidum'))

@app.route('/import_tahapan/<tahapan>', methods=['GET', 'POST'])
@login_required
def import_tahapan(tahapan):
    """Route untuk import data berdasarkan tahapan penanganan perkara"""
    # Redirect pra_penuntutan ke API khusus
    if tahapan == 'pra_penuntutan':
        return redirect(url_for('import_pra_penuntutan_api'))
    
    # Redirect upaya_hukum ke API khusus
    if tahapan == 'upaya_hukum':
        return redirect(url_for('import_upaya_hukum_api'))
    
    # Map tahapan URL ke tahapan database
    tahapan_mapping = {
        'penuntutan': 'PENUNTUTAN'
    }
    
    if tahapan not in tahapan_mapping:
        flash('Tahapan tidak valid', 'error')
        return redirect(url_for('input_pidum'))
    
    tahapan_penanganan = tahapan_mapping[tahapan]
    
    if request.method == 'POST':
        # Check if file is uploaded
        if 'file' not in request.files:
            flash('Tidak ada file yang dipilih', 'error')
            return redirect(request.url)
        
        file = request.files['file']
        if file.filename == '':
            flash('Tidak ada file yang dipilih', 'error')
            return redirect(request.url)
        
        if file:
            # Process the uploaded file with tahapan
            result = process_import_file(file, tahapan_penanganan)
            
            if result['success']:
                # Store import data in session for preview
                session['import_data'] = result['data']
                session['import_filename'] = file.filename
                session['import_tahapan'] = tahapan_penanganan
                
                # Add suggestions for jenis perkara
                for i, row in enumerate(result['data']):
                    original_jenis = row.get('JENIS_PERKARA_ORIGINAL', '')
                    suggestion = get_jenis_perkara_suggestions(original_jenis)
                    result['data'][i]['SUGGESTED_JENIS_PERKARA'] = suggestion
                
                return render_template('import_tahapan_preview.html', 
                                     import_data=result['data'],
                                     filename=file.filename,
                                     total_rows=result['total_rows'],
                                     tahapan_penanganan=tahapan_penanganan)
            else:
                flash(f'Error memproses file: {result["error"]}', 'error')
                return redirect(request.url)
    
    # Render template berdasarkan tahapan
    template_mapping = {
        'pra_penuntutan': 'import_pra_penuntutan.html',
        'penuntutan': 'import_penuntutan.html',
        'upaya_hukum': 'import_upaya_hukum.html'
    }
    
    return render_template(template_mapping[tahapan])

@app.route('/confirm_import_tahapan', methods=['POST'])
@login_required
def confirm_import_tahapan():
    """Confirm and process import with selected jenis perkara for specific tahapan"""
    import_data = session.get('import_data', [])
    tahapan_penanganan = session.get('import_tahapan', 'PRA PENUNTUTAN')
    
    if not import_data:
        flash('Tidak ada data import yang tersedia', 'error')
        return redirect(url_for('input_pidum'))
    
    try:
        # Get data from form
        prepared_data = []
        
        for i, original_row in enumerate(import_data):
            # Check if this row should be included (not removed)
            jenis_perkara_key = f'jenis_perkara_{i}'
            if jenis_perkara_key not in request.form:
                continue  # Skip removed rows
            
            # Get form data for this row
            periode = request.form.get(f'periode_{i}', '1')
            tanggal = request.form.get(f'tanggal_{i}', datetime.now().strftime('%Y-%m-%d'))
            jenis_perkara = request.form.get(jenis_perkara_key, 'PERKARA LAINNYA')
            keterangan = request.form.get(f'keterangan_{i}', '')
            
            # Prepare data for database insertion
            prepared_row = {
                'NO': str(original_row.get('NO', i + 1)),
                'PERIODE': str(periode),
                'TANGGAL': tanggal,
                'JENIS PERKARA': jenis_perkara,
                'TAHAPAN_PENANGANAN': tahapan_penanganan,
                'KETERANGAN': str(keterangan)
            }
            
            prepared_data.append(prepared_row)
        
        # Insert data to database
        success_count = 0
        error_count = 0
        error_details = []
        
        for data in prepared_data:
            try:
                insert_pidum_data(data)
                success_count += 1
            except Exception as e:
                error_count += 1
                error_details.append(f"Baris {data.get('NO', '?')}: {str(e)}")
                print(f"Error inserting row {data.get('NO', '?')}: {e}")
        
        # Clear session data
        session.pop('import_data', None)
        session.pop('import_filename', None)
        session.pop('import_tahapan', None)
        
        # Show results
        if success_count > 0:
            flash(f'Berhasil import {success_count} data {tahapan_penanganan}', 'success')
            if error_count > 0:
                flash(f'{error_count} data gagal diimport', 'warning')
                # Optionally show error details
                for error in error_details[:5]:  # Show max 5 errors
                    flash(error, 'warning')
        else:
            flash('Tidak ada data yang berhasil diimport', 'error')
            if error_details:
                for error in error_details[:3]:  # Show max 3 errors
                    flash(error, 'error')
            
    except Exception as e:
        flash(f'Error saat import data: {str(e)}', 'error')
    
    return redirect(url_for('view_pidum'))

@app.route('/import_pidum', methods=['GET', 'POST'])
@login_required
def import_pidum():
    """Route untuk import data PIDUM dari file Excel/CSV"""
    if request.method == 'POST':
        # Check if file is uploaded
        if 'file' not in request.files:
            flash('Tidak ada file yang dipilih', 'error')
            return redirect(request.url)
        
        file = request.files['file']
        if file.filename == '':
            flash('Tidak ada file yang dipilih', 'error')
            return redirect(request.url)
        
        if file:
            # Process the uploaded file
            result = process_import_file(file)
            
            if result['success']:
                # Store import data in session for preview
                session['import_data'] = result['data']
                session['import_filename'] = file.filename
                
                # Add suggestions for jenis perkara
                for i, row in enumerate(result['data']):
                    original_jenis = row.get('JENIS_PERKARA_ORIGINAL', '')
                    suggestion = get_jenis_perkara_suggestions(original_jenis)
                    result['data'][i]['SUGGESTED_JENIS_PERKARA'] = suggestion
                
                return render_template('import_preview.html', 
                                     import_data=result['data'],
                                     filename=file.filename,
                                     total_rows=result['total_rows'],
                                     data_type='PIDUM')
            else:
                flash(f'Error memproses file: {result["error"]}', 'error')
                return redirect(request.url)
    
    return render_template('import_pidum.html')

@app.route('/confirm_import_pidum', methods=['POST'])
@login_required
def confirm_import_pidum():
    """Confirm and process PIDUM import with selected jenis perkara"""
    import_data = session.get('import_data', [])
    
    if not import_data:
        flash('Tidak ada data import yang tersedia', 'error')
        return redirect(url_for('import_pidum'))
    
    try:
        # Get data from form
        prepared_data = []
        
        for i, original_row in enumerate(import_data):
            # Check if this row should be included (not removed)
            jenis_perkara_key = f'jenis_perkara_{i}'
            if jenis_perkara_key not in request.form:
                continue  # Skip removed rows
            
            # Get form data for this row
            periode = request.form.get(f'periode_{i}', '1')
            tanggal = request.form.get(f'tanggal_{i}', datetime.now().strftime('%Y-%m-%d'))
            jenis_perkara = request.form.get(jenis_perkara_key, 'PERKARA LAINNYA')
            
            # Prepare data for database insertion
            prepared_row = {
                'NO': str(original_row.get('NO', i + 1)),
                'PERIODE': str(periode),
                'TANGGAL': tanggal,
                'JENIS PERKARA': jenis_perkara
            }
            
            prepared_data.append(prepared_row)
        
        # Insert data to database
        success_count = 0
        error_count = 0
        error_details = []
        
        for data in prepared_data:
            try:
                insert_pidum_data(data)
                success_count += 1
            except Exception as e:
                error_count += 1
                error_details.append(f"Baris {data.get('NO', '?')}: {str(e)}")
                print(f"Error inserting row {data.get('NO', '?')}: {e}")
        
        # Clear session data
        session.pop('import_data', None)
        session.pop('import_filename', None)
        
        # Show results
        if success_count > 0:
            flash(f'Berhasil import {success_count} data PIDUM', 'success')
            if error_count > 0:
                flash(f'{error_count} data gagal diimport', 'warning')
                # Optionally show error details
                for error in error_details[:5]:  # Show max 5 errors
                    flash(error, 'warning')
        else:
            flash('Tidak ada data yang berhasil diimport', 'error')
            if error_details:
                for error in error_details[:3]:  # Show max 3 errors
                    flash(error, 'error')
            
    except Exception as e:
        flash(f'Error saat import data: {str(e)}', 'error')
    
    return redirect(url_for('view_pidum'))

# Register PDF conversion routes
try:
    from pdf_routes import register_pdf_routes
    register_pdf_routes(app)
    print("PDF conversion routes registered successfully")
except ImportError as e:
    print(f"Warning: Could not register PDF routes: {e}")

if __name__ == '__main__':
    print(f"Database initialized at: {get_database_stats()['database_path']}")
    app.run(debug=False, host='127.0.0.1', port=5002, threaded=True, use_reloader=False)